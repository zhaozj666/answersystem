from __future__ import annotations

from dataclasses import dataclass
from pathlib import Path
from typing import List

from docx import Document
from pypdf import PdfReader


@dataclass
class LoadedDocument:
    file_path: Path
    text: str


class DocumentLoader:
    def __init__(self, docs_dir: Path):
        self.docs_dir = docs_dir
        self.supported_suffixes = {".pdf", ".docx", ".txt", ".md"}

    def scan_files(self) -> List[Path]:
        files = [
            path
            for path in self.docs_dir.rglob("*")
            if path.is_file() and path.suffix.lower() in self.supported_suffixes
        ]
        files.sort()
        print(f"[INDEX] 扫描到文件: {[file.name for file in files]}")
        return files

    def load_file(self, file_path: Path) -> LoadedDocument:
        suffix = file_path.suffix.lower()
        if suffix in {".txt", ".md"}:
            text = file_path.read_text(encoding="utf-8", errors="ignore")
        elif suffix == ".docx":
            doc = Document(str(file_path))
            text = "\n".join(paragraph.text for paragraph in doc.paragraphs if paragraph.text.strip())
        elif suffix == ".pdf":
            reader = PdfReader(str(file_path))
            text = "\n".join(page.extract_text() or "" for page in reader.pages)
        else:
            raise ValueError(f"不支持的文件格式: {suffix}")

        print(f"[INDEX] 读取文件: {file_path.name}, 文本长度={len(text)}")
        return LoadedDocument(file_path=file_path, text=text)
